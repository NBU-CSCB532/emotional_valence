from newspaper import Article
from docx import Document
from openpyxl import load_workbook
from theguardian import theguardian_content
from datetime import date
import openpyxl
import newspaper
import requests
import datetime
import re
import pandas as pd
import os


class NewsDownloader:
    articles_downloaded_count = 0
    article_save_count = 1
    key_word = ""
    key_word_emotion = ""
    path = "Texts as found input/KeyWords.xlsx"
    article_last_link_list = [' '] * 1000
    article_all_downloaded_titles_set = set()
    maxArticlePerWord = 7

    def __init__(self):
        start_news_download_engine(self)
        try:
            cleanFromRepeatingArticle()
        except:
            print("Error during duplicates clean up.File sample.xlsx won't be created.")


def scrape_theguardian(keyword,date_input):
    """Scrapes theguardian world news through API request with given keyword and publication date"""
    query_keyword = '\"' + keyword + '\"'
    year, month, day = map(int, date_input.split('-'))
    date_to_compare = datetime.date(year, month, day)
    headers = {
        "q": query_keyword,
        "query-fields": "headline",
        "order-by": "relevance",
        "lang": "en",
        "section": "world",
        "show-fields": "headline,short-url, wordcount",
        "from-date" : date_to_compare
    }
    content = theguardian_content.Content(api='23c07455-eebd-4722-850f-acbf5a6fec2f', **headers)

    res = content.get_content_response()
    result = content.get_results(res)

    # print("Result {current_time}: {result}" .format(result=result, current_time = current_time))

    list_links = []
    for article in result:
        link = article.get('webUrl')
        list_links.append(link)

    return list_links


def valid_pubdate_input(string_input_date):
    """ Checks if user input is a valid date"""
    todays_date = date.today()
    year = int(string_input_date[0:4])
    if year<1950 or year> todays_date.year:
        return False

    month = int(string_input_date[5:7])
    if month<1 or month>12:
        return False

    day= int(string_input_date[8:10])
    if day<1 or day>31:
        return False

    return True

def start_news_download_engine(self):
    #asking user to input publication date from when on to search articles:
    date_input = input('Please enter a date from which you want to start downloading in YYYY-MM-DD format:')
    if not valid_pubdate_input(date_input):
        print("Date input is not valid. Date is automatically set to 2000-01-01")
        date_input = '2000-01-01'

    wb_obj = openpyxl.load_workbook(self.path)
    for s in range(0, 2):
        sheet_obj = wb_obj.worksheets[s]
        self.key_word_emotion = wb_obj.sheetnames[s].split()[1]

        print("Started downloading articles for " + self.key_word_emotion + "!\n")
        for x in range(2, sheet_obj.max_row):
            # Get current key word from keywords dictionary
            cell_obj = sheet_obj.cell(row=x, column=1)
            self.key_word = cell_obj.value
            if self.key_word is None:
                break
            print("Searching with key word:", self.key_word)

            # CNN Build
            print("STARTED CNN")
            res = requests.get("https://search.api.cnn.io/content?size=100&q=" + self.key_word + "&sort=relevance")
            links = [x['url'] for x in res.json()['result']]
            if len(links) < 6:
                continue
            init_article_download(self, links, "CNN", x, date_input)

            #TheGuardian build
            print("STARTED THEGUARDIAN")
            try:
                links_guardian = scrape_theguardian(self.key_word, date_input)
                init_article_download(self, links_guardian, "TheGuardian", x + 1, date_input)
            except:
                print(f"Error scraping the guardian for kew word: {self.key_word}")




def init_article_download(self, links_list, news_source, x, date_input):
    i = 0
    self.articles_downloaded_count += 1
    print(links_list)

    while self.articles_downloaded_count % self.maxArticlePerWord != 0 and i < (len(links_list) - 1):
        if self.article_last_link_list[x] != links_list[i]:
            #print(links_list)
            article = Article(links_list[i])
            news_get_text(self, article, news_source,date_input)
        else:
            print('Article already existing!')
            break;

        if i == 0:
            self.article_last_link_list[x] = links_list[0]
        i += 1

def remove_link(txt):
    for char_index in range(0, len(txt)):
        if txt[char_index].isupper() and txt[char_index + 1].isupper():
            end_index = char_index
            start_index = char_index
            while txt[char_index + 1].isupper() or txt[char_index + 1].isalpha() == 0 and len(txt) - char_index > 2:
                char_index += 1
                end_index += 1
            if end_index - start_index >= 35:
                txt = txt[0: start_index:] + txt[end_index - 1:len(txt)]

        else:
            char_index += 1
        if char_index == len(txt):
            return txt


def parse_and_compare_pbdates(string_pbdate, condition_date):
    """Parses publication date from article metadata and compares if it's after given date"""
    year = string_pbdate[0:4]
    month = string_pbdate[5:7]
    day= string_pbdate[8:10]
    #print(f'parsed datetime: {year}|{month}|{day}')
    datetime_pbdate = datetime.date(int(year),int(month),int(day))
    if datetime_pbdate > condition_date:
        return True
    else:
        return False

def news_get_text(self, article, news_source,date_input):
    article.download()
    try:
        article.parse()
    except:
        print("Not a valid URL: " + article.url)
        return 0

    #only take articles with publication date after provided date
    year, month, day = map(int, date_input.split('-'))
    date_to_compare = datetime.date(year, month, day)

    if news_source == "CNN":
        if article.url and article.publish_date is not None:
            article_meta_data = article.meta_data
            article_published_date = sorted({value for (key, value) in article_meta_data.items() if key == 'pubdate'})
            if article_published_date:
                is_after_date = parse_and_compare_pbdates(article_published_date[0],date_to_compare)
                if(is_after_date==False):
                    #print(f"Publication older than given for article '{article.title}'")
                    return 0
            else:
                #print(f'No published date for {article.title}')
                return 0
        else:
            return 0

    text_length = len(article.text.split())
    if text_length < 200 or text_length > 3000:
        #print(f'Article text length is {text_length} and not enough!')
        return 0

    if not (article.title in self.article_all_downloaded_titles_set):
        self.article_all_downloaded_titles_set.add(article.title)
    else:
        return 0

    # TO DO  remove all that don't match in text Title
    if self.key_word not in article.title:
        #print(f'Keyword "{self.key_word}" not in article title "{article.title}"')
        return 0

    self.articles_downloaded_count += 1
    try:
        parsed_news_text = remove_link(article.text)
    except:
        parsed_news_text = article.text

    save_to_xlsx_file(self, article, news_source, article.title)
    text_to_docx(self, parsed_news_text, article.title)
    print(f'Article {article.title} saved to docx from source {news_source}')


def save_to_xlsx_file(self, article, news_source, article_heading):
    wb = load_workbook("Texts as found input/" + "tmpSample.xlsx")
    ws = wb.active
    print(f"Number of saved articles = {self.article_save_count}")
    self.article_save_count += 1
    first_chars = article.title[0:16]

    ws["A" + str(self.article_save_count)] = first_chars
    ws['B' + str(self.article_save_count)] = article.title
    try:
        ws['C' + str(self.article_save_count)] = article.authors[0]
    except IndexError:
        ws['C' + str(self.article_save_count)] = "No author"

    ws['D' + str(self.article_save_count)] = news_source
    ws['E' + str(self.article_save_count)] = self.key_word
    ws['F' + str(self.article_save_count)] = self.key_word_emotion
    ws['G' + str(self.article_save_count)] = datetime.datetime.now()
    ws['H' + str(self.article_save_count)] = article.url
    ws['I' + str(self.article_save_count)] = "automatic search in the title"
    ws['J'+ str(self.article_save_count)] = re.sub('[^A-Za-z0-9]+', '', article_heading[0:16])

    wb.save("Texts as found input/" + "tmpSample.xlsx")


def text_to_docx(self, article_text, article_heading):
    document = Document()
    document.add_heading(article_heading, 0)
    p = document.add_paragraph(article_text)
    first_chars = re.sub('[^A-Za-z0-9]+', '', article_heading[0:16])

    test = 'Texts as found input/' + self.key_word_emotion + '/' + first_chars.replace(" ", "") + '.docx'
    document.save(test)


def findDuplicatesFilenames(data, sheet):
    """Finds duplicated articles in sample table and returns list with their associated docx filenames"""
    duplicated = data.duplicated(subset=['http'])
    duplicated_list = duplicated.values.tolist()
    row_ctr= 0
    filenames_to_delete = []
    for row in duplicated_list:
        row_ctr +=1
        if row :
            filename= sheet.cell(row=row_ctr+1, column=10).value
            filename_to_push = str(filename)+ '.docx'
            filenames_to_delete.append(filename_to_push)

    print(f"Filenames to delete are {filenames_to_delete}")
    return filenames_to_delete;

def deleteIfPathExists(filepath):
    """Deletes a file with the given filepath if it exists"""
    if os.path.exists(filepath):
        os.remove(filepath)
        print(f"Succesfully deleted {filepath}")
    else:
        print(f"The filepath {filepath} does not exist")

def deleteFiles(filenames_to_delete):
    """Deletes duplicated docx files in negate/positive subdirectories"""
    positive_downloaded_articles_path= "Texts as found input/" + "positive"
    negative_downloaded_articles_path = "Texts as found input/" + "negative"
    for filename in filenames_to_delete:
        filename_in_positive = positive_downloaded_articles_path + "/" + filename
        filename_in_negative = negative_downloaded_articles_path + "/" + filename
        deleteIfPathExists(filename_in_negative)
        deleteIfPathExists(filename_in_positive)


def cleanFromRepeatingArticle():
    """Removes duplicate articles in temporary sample excel table, deletes duplicated docx files and returns final cleaned sample table"""
    #print("Starting cleaning repeating articles and files.")
    sample_articles_excel_path = "Texts as found input/" + "tmpSample.xlsx"

    # loads panda dataframe from tmpSample excel table
    try:
        wb = openpyxl.load_workbook(sample_articles_excel_path)
        sheet = wb.active
        data= pd.read_excel(sample_articles_excel_path, keep_default_na=False)
        data= data.dropna()
    except:
        print(f"Troble in reading excel data from {sample_articles_excel_path}")
        return 0

    #finds and deletes duplicate files in folders
    try:
        filenames_to_delete = findDuplicatesFilenames(data,sheet);
        deleteFiles(filenames_to_delete)
    except:
        print(f"Error while trying to delete duplicated files.")

    #deletes duplicates from dataframe and saves cleaned data to sample excel table
    try:
        data.drop_duplicates(subset ="http",keep = False, inplace = True)
        final_sample_filename = "Texts as found input/" + "sample.xlsx"
        data.to_excel(final_sample_filename)
        wb_to_save = openpyxl.load_workbook(final_sample_filename)
        sheet_to_build = wb_to_save.active
        sheet_to_build.delete_cols(idx=1)
        wb_to_save.save(final_sample_filename)
    except:
        print(f"Trouble creating cleaned from duplicates sample excel table.")

