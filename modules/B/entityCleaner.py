import string

import docx
import spacy
import os
import re
from openpyxl import Workbook
import pandas as pd


class EntityCleaner:
    keyWords = ["positive", "negative"]

    my_path = os.path.abspath(os.path.dirname(__file__))
    input_path = None
    output_path = os.path.abspath(os.path.join(my_path, "TextsFiltered–protected"))
    all_stop_words = set()

    # A list that will be used in the checkLabel function.
    specificLabels = ["TIME", "DATE", "CARDINAL", "QUANTITY", "ORDINAL"]

    # Iterate through files in the directory.
    def __init__(self):
        for word in self.keyWords:
            input_path = os.path.abspath(os.path.join(self.my_path, "..\\..\Texts as found input\\" + word))
            textFiles = os.listdir(input_path)
            for file in textFiles:
                if file.find(".docx") == -1:
                    continue
                fileName = str(os.path.join(input_path, file))
                text = getText(self, fileName)
                text = editStopWords(self, text, file)
                saveEditedText(self, text, file)
                text2 = getText(self, fileName)
                saveStopWords(self, text2, file)

            print("All Texts are done and the Stop Words are added to the dictionary")

        merge_stop_words_files(self)


# Check if the format of the stopWord is correct. Returns true if the word should be classified as a stop word.
def checkLabel(self, stopWord, label):
    if label in self.specificLabels:
        # return bool(stopWord[0].isdigit()) # Check if the first character is a digit.
        return any(char.isdigit() for char in stopWord)  # Check if there are any digits in the stop word.
    return bool(1)


def checkWordLabel(self, stopWord, label):
    if label in self.specificLabels:
        # return bool(stopWord[0].isdigit()) # Check if the first character is a digit.
        return any(char.isdigit() for char in stopWord)  # Check if there are any digits in the stop word.
    return bool(1)


# Get text from .docx file and put it into variable "text".
def getText(self, fileName):
    doc = docx.Document(fileName)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)

    # Append new line after each paragraph
    return '\n'.join(text)


# Find stop words with Spacy's nlp, remove whitespace in them and add a tilde in the beginning of each of them.
def editStopWords(self, text, fileName):
    nlp = spacy.load("en_core_web_sm")
    doc = nlp(text)
    # A set that holds all stop words already replaced in text.
    stopWordsSet = set()

    # Iterate through all stop words, found by Spacy's nlp.
    for ent in doc.ents:
        # If the stop word is in the set, it has been already edited and replaced.
        if ent.text in stopWordsSet:
            continue

        if checkLabel(self, ent.text, ent.label_) == bool(0):
            continue

        # Create a copy of the stop word.
        stopWordReplace = ent.text
        stopWord = ent.text
        stopWord.replace(' \' s', '').replace('’s', '').replace('’', '').replace('“', '').replace('”', '').replace('”','').replace('”', '')
        stopWord = stopWord.translate(str.maketrans('', '', string.punctuation))



        # The following variable will hold the word that should be replaced. It is needed in case something should be removed from the string.

        # Check the stop word for "the" and remove it.
        if stopWord.find("the ") != -1:
            stopWord = stopWord.replace("the ", "")
            stopWordReplace = stopWordReplace.replace("the ", "")  # Modify the string that will be replaced

        # Check the stop word for a whitespace.
        if stopWord.find(" ") != -1:
            stopWord = stopWord.replace(" ", "")

        if stopWord.find("\n") != -1:
            stopWord = stopWord.replace("\n", "~\n")
        else:

            stopWord += "~"


        stopWordsSet.add(ent.text)

        # Replace the stop word with the edited copy.
        text = text.replace(stopWordReplace, stopWord)

    # Replace multiple tilde's at the end of some stop words.
    text = re.sub(r'~{2,}', "~", text)

    return text


# Recreate the file by adding paragraphs from "text" variable. Iterate through text and on every "\n" symbol create a new paragraph.
def saveEditedText(self, text, fileName):
    paragraph = ""
    newDoc = docx.Document()
    for character in text:
        if character == "\n":
            newDoc.add_paragraph(paragraph)
            paragraph = ""
        paragraph += character
    newDoc.save(
        os.path.abspath(
            os.path.join(os.path.abspath(os.path.dirname(__file__)), "..\\..\\TextsFiltered–protected", fileName)))


def saveStopWords(self, text, fileName):
    workbook = Workbook()
    worksheet = workbook.active

    worksheet.cell(row=1, column=1).value = 'Word'
    worksheet.cell(row=1, column=2).value = 'Comment'
    worksheet.cell(row=1, column=3).value = 'Stop word'
    worksheet.cell(row=1, column=4).value = 'Document'
    nlp = spacy.load("en_core_web_sm")
    doc = nlp(text)

    #
    # Addition of the stop words to the dictionary should be done here.
    #
    row = 2
    for ent in doc.ents:
        variable = ent.text


        if checkWordLabel(self, ent.text, ent.label_) == bool(1):
            if variable.find("the ") != -1:
                variable = variable.replace("the ", "")
            if variable.find(" ") != -1:
                variable = variable.replace(" ", "")
            if (variable.lower() in self.all_stop_words):
                continue
            else:
                variable.replace(' \' s', '').replace('’s', '').replace('’', '').replace('“', '').replace('”','').replace('”', '').replace('”', '')
                variable = variable.translate(str.maketrans('', '', string.punctuation))

                self.all_stop_words.add(variable.lower())
                worksheet.cell(row, 1).value = variable.lower() + "~"
                worksheet.cell(row, 2).value = ent.label_
                worksheet.cell(row, 3).value = 'Stop word'
                worksheet.cell(row, 4).value = fileName
                row += 1

    workbook.save(filename=os.path.abspath(
        os.path.join(os.path.abspath(os.path.dirname(__file__)), "..\\..\\Stop words", fileName[:-5]) + '.xlsx'))


def merge_stop_words_files(self):
    stop_words_path = os.path.abspath(os.path.join(self.my_path, "..\\..\\Stop words"))
    stop_words_excel_path = os.path.join(stop_words_path, 'Stop_words.xlsx')
    old_stop_words_path = os.path.abspath(os.path.join(self.my_path, "..\\..\\Dictionary"))
    old_stop_words = os.path.join(old_stop_words_path, 'Dictionary Stop Words.xlsx')
    df = pd.DataFrame()
    for f in os.listdir(stop_words_path):
        if f.endswith('.xlsx'):
            df = df.append(pd.read_excel(os.path.join(stop_words_path, f)), ignore_index=True)
    df.to_excel(stop_words_excel_path, index=False)
    data = pd.read_excel(stop_words_excel_path)
    data = data.sort_values("Word", ascending=True)
    data = data.drop_duplicates(subset="Word", keep='first', inplace=False)
    data.drop('Document', axis=1, inplace=True)
    data = data.rename(columns={'Word': ' Stop Words', 'Comment': 'comment', 'Stop word': 'is a stop word'})

    osw = pd.read_excel(old_stop_words)
    data = data.append(osw, ignore_index=True)
    data = data.sort_values(" Stop Words", ascending=True)
    data = data.drop_duplicates(subset=" Stop Words", keep='first', inplace=False)



    data.to_excel(os.path.abspath(
        os.path.join(os.path.abspath(os.path.dirname(__file__)), "..\\..\\Dictionary",
                     'Dictionary Stop Words.xlsx')),
        index=False)
