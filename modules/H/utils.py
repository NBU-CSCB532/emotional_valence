import unicodedata
import re
import os
import contextlib

from datetime import datetime
from docx import Document

from database import db_utils

from .common import SEARCH_TYPE_NEWS
from .common import SEARCH_TYPE_TWITTER
from .common import SEARCH_STATUS_STARTED


def slugify(value, allow_unicode=False):
    """
    Taken from https://github.com/django/django/blob/master/django/utils/text.py
    Convert to ASCII if 'allow_unicode' is False. Convert spaces or repeated
    dashes to single dashes. Remove characters that aren't alphanumerics,
    underscores, or hyphens. Convert to lowercase. Also strip leading and
    trailing whitespace, dashes, and underscores.
    """
    value = str(value)
    if allow_unicode:
        value = unicodedata.normalize('NFKC', value)
    else:
        value = unicodedata.normalize('NFKD', value).encode('ascii', 'ignore').decode('ascii')
    value = re.sub(r'[^\w\s-]', '', value.lower())
    return re.sub(r'[-\s]+', '-', value).strip('-_')


def get_doc_filename(title):
    return slugify(title) + '.docx'


def get_tweets_doc_filename(search_id, query):
    return 'twitter-search-{}-{}'.format(search_id, get_doc_filename(query))


def get_doc_filepath(filename, score):
    script_path = os.path.dirname(os.path.realpath(__file__))
    emotion = 'negative' if score < 0 else 'positive'
    return os.path.join(script_path, '..', '..', 'Texts as found input', emotion, filename)


def save_article_file(article, score):
    document = Document()
    document.add_heading(article.title, 0)
    document.add_paragraph(article.text)
    filename = get_doc_filename(article.title)
    filepath = get_doc_filepath(filename, score)
    document.save(filepath)


def save_tweets_file(search_id, query, tweets, score):
    document = Document()
    document.add_heading(query, 0)
    for tweet in tweets:
        document.add_paragraph(tweet.text)
    filename = get_tweets_doc_filename(search_id, query)
    filepath = get_doc_filepath(filename, score)
    document.save(filepath)


def execute_statement(statement, args=tuple()):
    with contextlib.closing(db_utils.db_connect()) as conn:
        with conn:
            with contextlib.closing(conn.cursor()) as cursor:
                cursor.execute(statement, args)


def select_many(statement, args=tuple()):
    with contextlib.closing(db_utils.db_connect()) as conn:
        with conn:
            with contextlib.closing(conn.cursor()) as cursor:
                return cursor.execute(statement, args).fetchall()


def select_one(statement, args=tuple()):
    with contextlib.closing(db_utils.db_connect()) as conn:
        with conn:
            with contextlib.closing(conn.cursor()) as cursor:
                return cursor.execute(statement, args).fetchone()


def save_news_search_to_db(query, articles, scores, from_date=None, until_date=None):
    with db_utils.db_connect() as conn:
        cursor = conn.cursor()
        result = cursor.execute("""
            INSERT INTO searches
            (timestamp, type, query, status, from_date, until_date)
            VALUES (?, ?, ?, ?, ?, ?)""",
            (datetime.now(), SEARCH_TYPE_NEWS, query, SEARCH_STATUS_STARTED, from_date, until_date))

        search_id = result.lastrowid
        conn.commit()

        data = [(
                    a.title,
                    a.url,
                    a.publish_date,
                    get_doc_filename(a.title),
                    search_id,
                    scores[a.url]
                ) for a in articles]

        cursor.executemany("""
            INSERT INTO documents
            (title, url, publish_date, filename, search_id, vader_score)
            VALUES (?, ?, ?, ?, ?, ?)""",
            data)
        conn.commit()

        return search_id


def save_twitter_search_to_db(query, tweets, scores):
    with db_utils.db_connect() as conn:
        cursor = conn.cursor()
        result = cursor.execute("""
            INSERT INTO searches
            (timestamp, type, query, status)
            VALUES (?, ?, ?, ?)""",
            (datetime.now(), SEARCH_TYPE_TWITTER, query, SEARCH_STATUS_STARTED))

        search_id = result.lastrowid
        conn.commit()

        filename = get_tweets_doc_filename(search_id, query)
        execute_statement(
                'UPDATE searches SET filename = ? WHERE id = ?',
                (filename, search_id))

        data = [(t.id, t.text, t.created_at, search_id, scores[t.id]) for t in tweets]
        cursor.executemany("""
            INSERT INTO tweets
            (tweet_id, text, created_at, search_id, vader_score)
            VALUES (?, ?, ?, ?, ?)""",
            data)
        conn.commit()

        return search_id


def get_documents_for_search(search_id):
    return select_many(
            'SELECT * FROM documents WHERE search_id = ?',
            (search_id,))


def update_document_biphone_score(document_id, score):
    execute_statement(
            'UPDATE documents SET biphone_score = ? WHERE id = ?',
            (score, document_id))


def update_search_status(search_id, status):
    execute_statement(
            'UPDATE searches SET status = ? WHERE id = ?',
            (status, search_id))


def update_search_biphone_score(search_id, score):
    execute_statement(
            'UPDATE searches SET biphone_score = ? WHERE id = ?',
            (score, search_id))


def get_searches():
    return select_many('SELECT * FROM searches ORDER BY timestamp DESC')


def get_search(id):
    return select_one('SELECT * FROM searches WHERE id = ?', (id,))


def get_tweets(search_id):
    return select_many(
            'SELECT * FROM tweets WHERE search_id = ? ORDER BY created_at DESC',
            (search_id,))


def read_document(title, score):
    filename = get_doc_filename(title)
    filepath = get_doc_filepath(filename, score)
    doc = Document(filepath)
    paragraphs = []
    for paragraph in doc.paragraphs:
        if paragraph:
            paragraphs += [p for p in paragraph.text.split('\n') if p != '']
    return paragraphs

